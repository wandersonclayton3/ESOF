import docx
doc = docx.Document('demo.docx')
len(doc.paragraphs)# Quantidade de Paragrafos no Documento
   
doc.paragraphs[0].text #Titulo do documento[linha 0 ]
   
doc.paragraphs[1].text #Texto do documento[linha 1 ]



doc.paragraphs[0].style = 'Normal' # Stilo da linha 1 = Normal
print(readDocx.getText('demo.docx'))#Printar na tela o texto do documento

doc.save('multipleParagraphs.docx')# Salvando o Documento com multipleParagraphs
